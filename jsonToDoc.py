# json_to_docx_converter.py

import json
import os
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

def aplicar_estilo_titulo(paragraph, nivel=1):
    """Aplica estilo personalizado aos títulos"""
    if nivel == 1:
        paragraph.style = 'Heading 1'
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.size = Pt(16)
        run.bold = True
    elif nivel == 2:
        paragraph.style = 'Heading 2'
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.size = Pt(14)
        run.bold = True

def processar_markdown_simples(doc, texto):
    """Processa markdown simples (## títulos e • bullets) e adiciona ao documento"""
    linhas = texto.split('\n')
    
    for linha in linhas:
        linha = linha.strip()
        if not linha:
            doc.add_paragraph('')  # Linha em branco
            continue
            
        if linha.startswith('## '):
            # Título nível 2
            titulo = linha[3:].strip()
            p = doc.add_paragraph(titulo)
            aplicar_estilo_titulo(p, nivel=2)
            
        elif linha.startswith('# '):
            # Título nível 1
            titulo = linha[2:].strip()
            p = doc.add_paragraph(titulo)
            aplicar_estilo_titulo(p, nivel=1)
            
        elif linha.startswith('• ') or linha.startswith('- '):
            # Bullet point
            texto_bullet = linha[2:].strip()
            p = doc.add_paragraph(texto_bullet, style='List Bullet')
            
        elif linha.startswith('**') and linha.endswith(':**'):
            # Texto em negrito com dois pontos
            titulo = linha[2:-3].strip()
            p = doc.add_paragraph()
            run = p.add_run(titulo + ':')
            run.bold = True
            
        else:
            # Texto normal
            doc.add_paragraph(linha)

def criar_docx_resumo_individual(arquivo_info, pasta_destino):
    """Cria um DOCX para um resumo individual"""
    
    doc = Document()
    
    # Título principal
    titulo_principal = f"Resumo - {arquivo_info.get('nome_original', arquivo_info['tipo'])}"
    p_titulo = doc.add_heading(titulo_principal, level=1)
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Informações do documento
    doc.add_heading('Informações do Documento', level=2)
    
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    
    # Preenche tabela de informações
    info_data = [
        ('Trimestre:', arquivo_info.get('trimestre', 'N/A')),
        ('Tipo:', arquivo_info.get('nome_original', arquivo_info['tipo'])),
        ('Arquivo:', arquivo_info.get('nome_arquivo', Path(arquivo_info['arquivo']).name)),
        ('Status:', arquivo_info['status'].title()),
        ('Data Processamento:', arquivo_info.get('timestamp', 'N/A')[:19])
    ]
    
    for i, (chave, valor) in enumerate(info_data):
        info_table.cell(i, 0).text = chave
        info_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        info_table.cell(i, 1).text = str(valor)
    
    doc.add_paragraph('')  # Espaço
    
    # Resumo do conteúdo
    if arquivo_info['status'] == 'sucesso' and arquivo_info.get('resumo'):
        doc.add_heading('Resumo Executivo', level=2)
        processar_markdown_simples(doc, arquivo_info['resumo'])
    else:
        doc.add_heading('Erro no Processamento', level=2)
        doc.add_paragraph(f"Erro: {arquivo_info.get('erro', 'Erro desconhecido')}")
    
    # Metadados adicionais (se disponível)
    if arquivo_info.get('num_paginas'):
        doc.add_paragraph('')
        doc.add_paragraph(f"Número de páginas processadas: {arquivo_info['num_paginas']}")
    
    if arquivo_info.get('tamanho_texto'):
        doc.add_paragraph(f"Tamanho do texto processado: {arquivo_info['tamanho_texto']:,} caracteres")
    
    # Nome do arquivo DOCX
    nome_arquivo = f"resumo_{arquivo_info['tipo']}_{arquivo_info.get('trimestre', 'sem_trimestre')}.docx"
    caminho_arquivo = pasta_destino / nome_arquivo
    
    # Salva o documento
    doc.save(str(caminho_arquivo))
    return caminho_arquivo

def criar_docx_resumo_executivo(dados_analise, pasta_destino):
    """Cria DOCX do resumo executivo consolidado"""
    
    doc = Document()
    
    # Título principal
    titulo = f"Resumo Executivo - {dados_analise['trimestre']}"
    p_titulo = doc.add_heading(titulo, level=1)
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Informações gerais
    doc.add_heading('Informações Gerais', level=2)
    
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ('Trimestre:', dados_analise['trimestre']),
        ('Data Processamento:', dados_analise.get('timestamp', 'N/A')[:19]),
        ('Arquivos Processados:', f"{len([a for a in dados_analise.get('arquivos_processados', []) if a['status'] == 'sucesso'])}/{len(dados_analise.get('arquivos_processados', []))}"),
        ('Status:', dados_analise.get('status', 'N/A').title())
    ]
    
    for i, (chave, valor) in enumerate(info_data):
        info_table.cell(i, 0).text = chave
        info_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        info_table.cell(i, 1).text = str(valor)
    
    doc.add_paragraph('')
    
    # Resumo Executivo
    if dados_analise.get('resumo_executivo'):
        doc.add_heading('Resumo Executivo Consolidado', level=2)
        processar_markdown_simples(doc, dados_analise['resumo_executivo'])
    
    # Lista de arquivos processados
    doc.add_heading('Arquivos Processados', level=2)
    
    for arquivo in dados_analise.get('arquivos_processados', []):
        status_icon = "✅" if arquivo['status'] == 'sucesso' else "❌"
        nome = arquivo.get('nome_original', arquivo['tipo'])
        doc.add_paragraph(f"{status_icon} {nome} - {arquivo['status'].title()}")
        
        if arquivo['status'] == 'erro':
            p_erro = doc.add_paragraph(f"    Erro: {arquivo.get('erro', 'N/A')}")
            p_erro.style = 'Intense Quote'
    
    # Nome do arquivo
    nome_arquivo = f"resumo_executivo_{dados_analise['trimestre']}.docx"
    caminho_arquivo = pasta_destino / nome_arquivo
    
    doc.save(str(caminho_arquivo))
    return caminho_arquivo

def converter_json_para_docx(caminho_json, pasta_destino=None):
    """
    Função principal que converte um JSON de análise para documentos DOCX
    
    Args:
        caminho_json (str): Caminho para o arquivo JSON
        pasta_destino (str): Pasta onde salvar os DOCX (opcional)
    
    Returns:
        dict: Resultado da conversão
    """
    
    caminho_json = Path(caminho_json)
    
    if not caminho_json.exists():
        return {
            'status': 'erro',
            'erro': f'Arquivo JSON não encontrado: {caminho_json}'
        }
    
    # Define pasta destino
    if pasta_destino:
        pasta_destino = Path(pasta_destino)
    else:
        pasta_destino = caminho_json.parent / f"documentos_{caminho_json.stem}"
    
    pasta_destino.mkdir(exist_ok=True)
    
    try:
        # Carrega dados do JSON
        with open(caminho_json, 'r', encoding='utf-8') as f:
            dados = json.load(f)
        
        arquivos_criados = []
        
        # Cria DOCX do resumo executivo se disponível
        if dados.get('resumo_executivo'):
            caminho_executivo = criar_docx_resumo_executivo(dados, pasta_destino)
            arquivos_criados.append({
                'tipo': 'resumo_executivo',
                'arquivo': str(caminho_executivo)
            })
            print(f"✅ Resumo executivo criado: {caminho_executivo.name}")
        
        # Cria DOCX para cada arquivo processado individualmente
        for arquivo_info in dados.get('arquivos_processados', []):
            if arquivo_info['status'] == 'sucesso':
                caminho_individual = criar_docx_resumo_individual(arquivo_info, pasta_destino)
                arquivos_criados.append({
                    'tipo': 'resumo_individual',
                    'nome_original': arquivo_info.get('nome_original', arquivo_info['tipo']),
                    'arquivo': str(caminho_individual)
                })
                print(f"✅ Resumo individual criado: {caminho_individual.name}")
        
        return {
            'status': 'sucesso',
            'pasta_destino': str(pasta_destino),
            'arquivos_criados': arquivos_criados,
            'total_arquivos': len(arquivos_criados)
        }
        
    except json.JSONDecodeError as e:
        return {
            'status': 'erro',
            'erro': f'Erro ao ler JSON: {e}'
        }
    except Exception as e:
        return {
            'status': 'erro',
            'erro': f'Erro inesperado: {e}'
        }

def processar_pasta_resultados(pasta_resultados="resultados_analises"):
    """Processa todos os JSONs de análise em uma pasta"""
    
    pasta = Path(pasta_resultados)
    if not pasta.exists():
        print(f"❌ Pasta não encontrada: {pasta}")
        return False
    
    jsons_analise = list(pasta.glob("analise_*.json"))
    
    if not jsons_analise:
        print(f"⚠️ Nenhum arquivo de análise encontrado em {pasta}")
        return False
    
    print(f"📁 Encontrados {len(jsons_analise)} arquivos de análise")
    
    for json_file in jsons_analise:
        print(f"\n🔄 Processando: {json_file.name}")
        resultado = converter_json_para_docx(str(json_file))
        
        if resultado['status'] == 'sucesso':
            print(f"✅ {resultado['total_arquivos']} documentos DOCX criados")
        else:
            print(f"❌ Erro: {resultado['erro']}")
    
    return True

# Função de teste
def testar_conversao():
    """Testa a conversão com um JSON de exemplo"""
    
    # Procura por arquivos JSON na pasta de resultados
    pasta_resultados = Path("resultados_analises")
    
    if pasta_resultados.exists():
        jsons = list(pasta_resultados.glob("analise_*.json"))
        if jsons:
            print(f"📋 Testando com: {jsons[0].name}")
            resultado = converter_json_para_docx(str(jsons[0]))
            print(f"Resultado: {resultado}")
        else:
            print("⚠️ Nenhum arquivo de análise encontrado para teste")
    else:
        print("⚠️ Pasta resultados_analises não encontrada")


    
    